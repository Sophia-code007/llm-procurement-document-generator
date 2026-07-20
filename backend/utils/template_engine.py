import re
from docx import Document


class BidTemplateEngine:
    """招标文件动态模板引擎"""

    def __init__(self, template_path):
        self.template_path = template_path
        self.doc = None

    def _replace_in_paragraph(self, paragraph, data):
        """在单个段落中替换占位符"""
        full_text = paragraph.text
        if '{{' not in full_text:
            return

        new_text = full_text
        for key, value in data.items():
            placeholder = '{{' + key + '}}'
            if placeholder in new_text:
                new_text = new_text.replace(placeholder, str(value))

        if new_text != full_text:
            for run in paragraph.runs:
                run.text = ''
            if paragraph.runs:
                paragraph.runs[0].text = new_text
            else:
                paragraph.add_run(new_text)

    def _replace_in_table(self, table, data):
        """在表格中替换占位符"""
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    self._replace_in_paragraph(paragraph, data)

    def render(self, data: dict, output_path: str):
        """
        渲染模板，生成最终文件
        :param data: 字段字典，key为占位符名称（不含{{}}），value为替换值
        :param output_path: 输出文件路径
        """
        self.doc = Document(self.template_path)

        # 替换正文段落
        for paragraph in self.doc.paragraphs:
            self._replace_in_paragraph(paragraph, data)

        # 替换所有表格
        for table in self.doc.tables:
            self._replace_in_table(table, data)

        # 替换页眉页脚
        for section in self.doc.sections:
            for paragraph in section.header.paragraphs:
                self._replace_in_paragraph(paragraph, data)
            for paragraph in section.footer.paragraphs:
                self._replace_in_paragraph(paragraph, data)

        self.doc.save(output_path)
        return output_path

    def list_placeholders(self) -> list:
        """列出模板中所有占位符"""
        doc = Document(self.template_path)
        placeholders = set()
        pattern = re.compile(r'\{\{(\w+)\}\}')

        # 扫描段落
        for p in doc.paragraphs:
            matches = pattern.findall(p.text)
            placeholders.update(matches)

        # 扫描表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        matches = pattern.findall(p.text)
                        placeholders.update(matches)

        return sorted(list(placeholders))


# ============================================================
# 字段映射表：招标模板占位符映射
# ============================================================
BID_TEMPLATE_FIELDS = {
    # ===== 封面及基础信息 =====
    'projectCode': {
        'label': '项目编号',
        'category': '基础信息',
        'example': 'GC-GG2026-001',
        'required': True
    },
    'projectName': {
        'label': '项目名称',
        'category': '基础信息',
        'example': 'XX办公楼装修改造工程',
        'required': True
    },
    'tendererName': {
        'label': '招标人名称',
        'category': '基础信息',
        'example': 'XX机关事务管理局',
        'required': True
    },
    'legalRepresentative': {
        'label': '法定代表人',
        'category': '基础信息',
        'example': '张三',
        'required': False
    },
    'issueYear': {
        'label': '发布年份',
        'category': '基础信息',
        'example': '2026',
        'required': True
    },
    'issueMonth': {
        'label': '发布月份',
        'category': '基础信息',
        'example': '7',
        'required': True
    },
    'issueDay': {
        'label': '发布日期',
        'category': '基础信息',
        'example': '15',
        'required': True
    },

    # ===== 项目概况（投标须知前附表） =====
    'constructionLocation': {
        'label': '建设地点',
        'category': '项目概况',
        'example': '北京市西城区XX大街XX号',
        'required': True
    },
    'constructionScale': {
        'label': '工程规模（平方米）',
        'category': '项目概况',
        'example': '5000',
        'required': True
    },
    'structureType': {
        'label': '结构形式',
        'category': '项目概况',
        'example': '框架结构',
        'required': True
    },
    'eavesHeight': {
        'label': '檐高（米）',
        'category': '项目概况',
        'example': '18',
        'required': False
    },
    'aboveGroundFloors': {
        'label': '地上层数',
        'category': '项目概况',
        'example': '5',
        'required': True
    },
    'undergroundFloors': {
        'label': '地下层数',
        'category': '项目概况',
        'example': '0',
        'required': True
    },
    'fundSource': {
        'label': '资金来源',
        'category': '项目概况',
        'example': '国拨，资金已到位',
        'required': True
    },

    # ===== 招标信息 =====
    'bidScope': {
        'label': '招标范围',
        'category': '招标信息',
        'example': '图纸所示全部建筑装修工程',
        'required': True
    },
    'constructionPeriod': {
        'label': '要求工期（日历天）',
        'category': '招标信息',
        'example': '120',
        'required': True
    },
    'planStartDate': {
        'label': '计划开工日期',
        'category': '招标信息',
        'example': '2026年8月1日',
        'required': True
    },
    'planEndDate': {
        'label': '计划竣工日期',
        'category': '招标信息',
        'example': '2026年11月28日',
        'required': True
    },
    'qualityStandard': {
        'label': '质量标准',
        'category': '招标信息',
        'example': '合格',
        'required': True
    },
    'qualificationRequirement': {
        'label': '投标人资质要求',
        'category': '招标信息',
        'example': '建筑装修装饰工程专业承包二级及以上',
        'required': True
    },

    # ===== 踏勘现场 =====
    'siteVisitDate': {
        'label': '踏勘日期',
        'category': '踏勘安排',
        'example': '2026年7月25日',
        'required': True
    },
    'siteVisitTime': {
        'label': '踏勘时间',
        'category': '踏勘安排',
        'example': '上午9:00时',
        'required': True
    },
    'siteVisitLocation': {
        'label': '集合地点',
        'category': '踏勘安排',
        'example': '项目现场大门口',
        'required': True
    },

    # ===== 投标相关 =====
    'bidValidPeriod': {
        'label': '投标有效期',
        'category': '投标相关',
        'example': '90日历日',
        'required': True
    },
    'bidBondAmount': {
        'label': '投标保证金金额（万元）',
        'category': '投标相关',
        'example': '0',
        'required': True
    },
    'docCopyCount': {
        'label': '投标文件份数',
        'category': '投标相关',
        'example': '一式两份，正本一份副本一份',
        'required': True
    },
    'drawingDeposit': {
        'label': '图纸押金（元）',
        'category': '投标相关',
        'example': '500',
        'required': False
    },

    # ===== 投标递交 =====
    'submitLocation': {
        'label': '递交地点',
        'category': '投标递交',
        'example': '中央国家机关政府采购中心开标室',
        'required': True
    },
    'submitDeadlineDate': {
        'label': '递交截止日期',
        'category': '投标递交',
        'example': '2026年8月5日',
        'required': True
    },
    'submitDeadlineTime': {
        'label': '递交截止时间',
        'category': '投标递交',
        'example': '上午9:30',
        'required': True
    },

    # ===== 开标 =====
    'bidOpenDate': {
        'label': '开标日期',
        'category': '开标安排',
        'example': '2026年8月5日',
        'required': True
    },
    'bidOpenTime': {
        'label': '开标时间',
        'category': '开标安排',
        'example': '上午9:30',
        'required': True
    },
    'bidOpenLocation': {
        'label': '开标地点',
        'category': '开标安排',
        'example': '中央国家机关政府采购中心开标室',
        'required': True
    },

    # ===== 评标 =====
    'evaluationMethod': {
        'label': '评标方法',
        'category': '评标',
        'example': '综合定量评标法',
        'required': True
    },
    'contractType': {
        'label': '合同形式',
        'category': '评标',
        'example': '固定总价合同',
        'required': True
    },
    'ceilingPrice': {
        'label': '拦标价（万元）',
        'category': '评标',
        'example': '500',
        'required': False
    },
    'reserveFund': {
        'label': '预留金（万元）',
        'category': '评标',
        'example': '50',
        'required': False
    },

    # ===== 招标人联系方式 =====
    'tendererAddress': {
        'label': '招标人地址',
        'category': '联系方式',
        'example': '北京市西城区XX大街XX号',
        'required': True
    },
    'contactPerson': {
        'label': '联系人',
        'category': '联系方式',
        'example': '李工',
        'required': True
    },
    'contactPhone': {
        'label': '联系电话',
        'category': '联系方式',
        'example': '010-XXXXXXX',
        'required': True
    },
}
